import requests
from bs4 import BeautifulSoup
from datetime import datetime
import pytz  
import os
import logging
from pathlib import Path
from docx import Document
import re


# Get the absolute path to the logs directory
LOG_DIR = Path(__file__).resolve().parent.parent.parent / "logs"
# Ensure the logs directory exists
LOG_DIR.mkdir(parents=True, exist_ok=True)  

# Define the log file path
LOG_FILE = LOG_DIR / "scraper.log"


# Configure logging
logging.basicConfig(
    filename=LOG_FILE,
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    level=logging.INFO
)

logging.info("Scraper started.")


# Define storage directory
STORAGE_DIR = Path(__file__).resolve().parent.parent.parent / "data"
STORAGE_DIR.mkdir(exist_ok=True)  # Ensure the directory exists


# Constants
BASE_URL = "https://techcrunch.com/latest/"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36",
    "Accept-Language": "en-US,en;q=0.5",
}

def fetch_page(url):
    """Fetches and parses a webpage, returning a BeautifulSoup object."""
    try:
        response = requests.get(url, headers=HEADERS, timeout=10)
        response.raise_for_status()
        return BeautifulSoup(response.text, "lxml")
    except requests.RequestException as e:
        print(f"Error fetching {url}: {e}")
        logging.error(f"Error fetching {url}: {e}")
        return None

def extract_article_links(soup, limit=3):
    """Extracts the latest article links."""
    if not soup:
        logging.warning("Soup object is None. No articles extracted....")
        return []

    links = []
    for article in soup.find_all("div", class_="wp-block-techcrunch-card", limit=limit):
        title_tag = article.find("h3", class_="loop-card__title")
        link_tag = title_tag.find("a") if title_tag else None
        if link_tag and link_tag["href"]:
            links.append(link_tag["href"])  # Extract article URLs

    if not links:
        logging.warning("No article links found.")

    return links

def extract_authors(soup):
    """Extracts the author(s) from the article's meta tag or author list."""
    meta_author = soup.find("meta", attrs={"name": "author"})
    if meta_author and meta_author["content"].strip():
        return meta_author["content"].strip()

    authors = set()
    for author_list in soup.find_all("ul", class_="post-authors-list__author-list"):
        for author_tag in author_list.find_all("a", class_="post-authors-list__author"):
            authors.add(author_tag.text.strip())

    return ", ".join(authors) if authors else "Unknown"

def extract_article_content(soup):
    """Extracts the main content of the article."""
    excerpt_tag = soup.find("p", class_="wp-block-techcrunch-storyline-hero__excerpt")
    excerpt = excerpt_tag.text.strip().replace("\n\n", " ") if excerpt_tag else ""

    paragraphs = soup.find_all("p", class_="wp-block-paragraph")
    # content = "\n".join(
    #     p.text.strip() for p in paragraphs 
    #     if p.text.strip() and not p.text.strip().startswith("Topics") and "© 2024 Yahoo" not in p.text.strip()
    # )
    content = "\n".join(
    p.text.strip() for p in paragraphs
    if p.text.strip() and not p.text.strip().startswith("Topics")
    )

    # Remove copyright patterns like "© 2025 Yahoo" or similar
    content = re.sub(r"© \d{4} Yahoo.", "", content)

    if not content:
        logging.warning("Article content is empty or unavailable....")

    return f"{excerpt}\n{content}".strip() if excerpt else content or "Content not available"

def format_published_time(time_str):
    """Converts an ISO 8601 datetime string into a human-readable format."""
    try:
        dt = datetime.fromisoformat(time_str)
        local_dt = dt.astimezone(pytz.timezone("Africa/Cairo"))
        return local_dt.strftime("%B %d, %Y at %I:%M %p (%Z)")
    except Exception:
        logging.warning(f"Invalid date format: {time_str}")
        return "Unknown Date"

def fetch_article_details(url):
    """Fetches full details (title, author, published time, content) from a given article URL."""
    soup = fetch_page(url)
    if not soup:
        logging.error(f"Failed to fetch article content from {url}")
        return {"url": url, "title": "Unknown", "authors": "Unknown", "published_time": "Unknown", "content": "Failed to fetch content"}

    title_tag = soup.find("title")
    title = title_tag.text.strip().replace(" | TechCrunch", "") if title_tag else "No Title"

    time_tag = soup.find("time")
    published_time = format_published_time(time_tag["datetime"]) if time_tag and time_tag.has_attr("datetime") else "Unknown"

    article_data = {
        "url": url,
        "title": title,
        "authors": extract_authors(soup),
        "published_time": published_time,
        "content": extract_article_content(soup),
    }

    logging.info(f"Successfully fetched article: {title}")

    return article_data

def get_latest_articles(limit=3):
    """Fetches the latest articles' details."""
    soup = fetch_page(BASE_URL)
    article_links = extract_article_links(soup, limit)

    if not article_links:
        print("No articles found.")
        logging.warning("No articles found.")
        return []

    articles = [fetch_article_details(url) for url in article_links]

    logging.info(f"Successfully fetched {len(articles)} articles.")

    return articles



def store_article(article_data, file_format="txt"):
    """
    Stores the article in the specified format (txt or docx).
    """
    # Clean title for filename
    title = article_data["title"].replace("/", "-")  
    filename = STORAGE_DIR / f"{title}.{file_format}"

    try:
        if file_format == "txt":
            with open(filename, "w", encoding="utf-8") as f:
                f.write(f"Title: {article_data['title']}\n")
                f.write(f"Author(s): {article_data['authors']}\n")
                f.write(f"Published: {article_data['published_time']}\n\n")
                f.write(article_data["content"])
        elif file_format == "docx":
            doc = Document()
            doc.add_heading(article_data["title"], level=1)
            doc.add_paragraph(f"Author(s): {article_data['authors']}")
            doc.add_paragraph(f"Published: {article_data['published_time']}")
            doc.add_paragraph("\n" + article_data["content"])
            doc.save(filename)
        else:
            logging.warning(f"Unsupported file format: {file_format}")
            return False
        
        logging.info(f"Article saved successfully: {filename}")
        return True
    except Exception as e:
        logging.error(f"Error saving article {article_data['title']}: {e}")
        return False



#For testing :
if __name__ == "__main__":
    articles = get_latest_articles(limit=3)

    for i, article in enumerate(articles, 1):

        store_article(article, "docx")
        store_article(article, "txt")
        print(f"{i}. {article['title']}")
        print(f"   Link: {article['url']}")
        print(f"   Published: {article['published_time']}")
        print(f"   Author(s): {article['authors']}")
        print(f"   Content Preview:\n{article['content']}")  
        print("-" * 80)

