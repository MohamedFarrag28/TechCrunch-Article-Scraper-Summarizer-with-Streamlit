import streamlit as st
import logging
import sys
import json
from pathlib import Path
import pandas as pd
import docx


# Configure logging
logging.basicConfig(
    filename="logs/app.log",
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    level=logging.INFO,
)


# Ensure project root is added to sys.path
ROOT_DIR = Path(__file__).resolve().parent.parent
sys.path.append(str(ROOT_DIR))

from src.scraping.scraper import get_latest_articles, fetch_article_details
from src.summarization.summarizer import summarize_text
from src.utils.helpers import *


st.set_page_config(page_title="TechCrunch Summarizer", layout="wide")

st.title("🚀 TechCrunch Article Summarizer")
st.sidebar.header("🔍 Options")

# --- Sidebar: Fetch Latest Articles ---
with st.sidebar.expander("📡 Fetch TechCrunch Articles", expanded=True):
    if st.button("Get Latest Articles"):
        with st.spinner("Fetching articles..."):
            articles = get_latest_articles(limit=5)  # Fetch 5 articles
            st.session_state["articles"] = articles
            st.session_state["selected_article"] = None
            st.session_state["summary"] = None  
            st.session_state["summary_stats"] = None
            st.success("Articles fetched successfully!")

# --- Sidebar: URL Input ---
with st.sidebar.expander("🔗 Summarize Custom Article", expanded=True):
    custom_url = st.text_input("Enter Article URL")
    if st.button("Fetch & Summarize"):
        if custom_url:
            with st.spinner("Fetching article..."):
                custom_article = fetch_article_details(custom_url)
                st.session_state["articles"] = [custom_article]
                st.session_state["selected_article"] = None
                st.session_state["summary"] = None
                st.session_state["summary_stats"] = None
                st.success("Article fetched successfully!")
        else:
            st.error("Please enter a valid URL.")


# --- Sidebar: Overall Feedback ---
st.sidebar.subheader("📝 Overall Feedback")

overall_feedback_text = st.sidebar.text_area("Your feedback here:")

if st.sidebar.button("Submit_Feedback"):
    validation_error = validate_text_input(overall_feedback_text)
    if validation_error:
        st.sidebar.error(validation_error)
    else:
        feedback_entry = {
            "type": "overall_feedback",
            "message": overall_feedback_text,
            "input": "General"  # No specific article, so we mark it as "General"
        }
        save_feedback(feedback_entry)
        st.sidebar.success("✅ Feedback submitted successfully!")

# --- Display Articles ---
if "articles" in st.session_state and st.session_state["articles"]:
    articles = st.session_state["articles"]
    article_titles = [f"{i+1}. {art['title']}" for i, art in enumerate(articles)]
    
    selected_article = st.selectbox("📑 Select an article:", article_titles, key="article_select")

    # Reset summary and statistics when a new article is selected
    if selected_article != st.session_state.get("selected_article"):
        st.session_state["selected_article"] = selected_article
        st.session_state["summary"] = None
        st.session_state["summary_stats"] = None

    if selected_article:
        article_index = article_titles.index(selected_article)
        article = articles[article_index]

        # --- Display Article ---
        st.subheader(f"📰 {article['title']}")
        st.write(f"**✍️ Author(s):** {article['authors']}")
        st.write(f"**🗓 Published:** {article['published_time']}")
        st.write(f"**🔗 [Read Full Article]({article['url']})**")
        st.markdown("---")
        
        st.subheader("📜 Full Article Content")
        st.text_area("Article Text", article["content"], height=400, disabled=True)


        st.subheader("⚙️ Summary Settings")
        summary_length = st.slider("Select summary length (words)", min_value=50, max_value=300, value=200, step=10)

        # --- Summarization Button ---
        if st.button("Summarize Article", key="summarize_button"):
            with st.spinner("Summarizing..."):
                summary, stats = summarize_text(article["content"], max_length=summary_length)

                # ✅ Store results in session state
                st.session_state["summary"] = summary
                st.session_state["summary_stats"] = stats  

                st.success("Summarization complete! ✅")

        # --- Display Summary ---
        if "summary" in st.session_state and st.session_state["summary"]:
            st.subheader("📝 AI-Generated Summary")
            st.text_area("Summary", st.session_state["summary"], height=100, disabled=False)

            # --- Save as DOCX ---
            doc = docx.Document()
            doc.add_heading("TechCrunch Article Summary", level=1)
            doc.add_paragraph(f"Title: {article['title']}")
            doc.add_paragraph(f"Author(s): {article['authors']}")
            doc.add_paragraph(f"Published: {article['published_time']}")
            doc.add_paragraph("\nSummary:\n")
            doc.add_paragraph(st.session_state["summary"])

            doc_path = "temp/summary.docx"
            doc.save(doc_path)

            with open(doc_path, "rb") as f:
                st.download_button("📥 Download DOCX", f, file_name="summary.docx")

            # --- Save as TXT ---
            txt_content = (
                f"Title: {article['title']}\n"
                f"Author(s): {article['authors']}\n"
                f"Published: {article['published_time']}\n\n"
                f"Summary:\n{st.session_state['summary']}"
            )

            txt_path = "temp/summary.txt"
            with open(txt_path, "w",encoding='utf-8') as f:
                f.write(txt_content)

            with open(txt_path, "rb") as f:
                st.download_button("📥 Download TXT", f, file_name="summary.txt")

            # Show Statistics Button
            if st.button("Show Statistics", key="stats_button"):
                stats = st.session_state["summary_stats"]
                with st.expander("📊 Summary Statistics", expanded=True):
                    st.write("🔢 **Compression Ratio:**", stats.get("compression_ratio", "N/A"))
                    st.write("📖 **readability_grade(flesch_kincaid):**", stats.get("readability_grade(flesch_kincaid)", "N/A"))
                    st.write("📖 **Reading Ease(fl_reading_ease):**", stats.get("Reading Ease(fl_reading_ease )", "N/A"))
                    st.write("📌 **Total Words in Article:**", stats.get("original_words", "N/A"))
                    st.write("✂️ **Total Words in Summary:**", stats.get("summary_words", "N/A"))


            st.markdown("---")  # Separator Line

            st.subheader("💡 Feedback")
            feedback_text = st.text_area("Let us know your thoughts about the summary:", key="feedback_input")

            if st.button("Submit Feedback"):
                validation_error = validate_text_input(feedback_text)
                if validation_error:
                    st.error(validation_error)
                else:
                    feedback_entry = {
                        "type": "summary_feedback",
                        "message": feedback_text,
                        "input": article["content"],             # Saving the article
                        "summary": st.session_state['summary']   # Saving the summary
                    }
                    save_feedback(feedback_entry)
                    st.success("✅ Feedback submitted successfully!")



    